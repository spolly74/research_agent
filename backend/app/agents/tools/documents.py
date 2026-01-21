"""
Document Parser Tools - Parse PDF and DOCX files.

This module provides tools for:
- Extracting text from PDF files
- Extracting text from DOCX files
- Document metadata extraction
"""

import os
import re
from pathlib import Path
from typing import Optional

import structlog
from langchain_core.tools import tool

logger = structlog.get_logger(__name__)

# Configuration
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB max
MAX_OUTPUT_LENGTH = 100000  # Max characters returned
MAX_PAGES = 500  # Max pages to process

# Allowed directories for document reading
ALLOWED_DOCUMENT_DIRS = [
    os.path.expanduser("~/research_data"),
    os.path.expanduser("~/Documents"),
    "/tmp/research_agent",
]

# Supported extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def _is_document_path_safe(path: str) -> tuple[bool, str]:
    """
    Validate that a document path is safe to access.

    Args:
        path: The file path to validate

    Returns:
        Tuple of (is_safe, error_message)
    """
    if ".." in path:
        return False, "Path traversal (..) not allowed"

    try:
        resolved = Path(path).resolve()
    except (ValueError, OSError) as e:
        return False, f"Invalid path: {e}"

    # Check if path is within allowed directories
    allowed = False
    for allowed_dir in ALLOWED_DOCUMENT_DIRS:
        try:
            allowed_path = Path(allowed_dir).resolve()
            if str(resolved).startswith(str(allowed_path)):
                allowed = True
                break
        except (ValueError, OSError):
            continue

    if not allowed:
        return False, f"Path must be within allowed directories: {ALLOWED_DOCUMENT_DIRS}"

    return True, ""


def configure_allowed_document_dirs(dirs: list[str]) -> None:
    """
    Configure allowed directories for document parsing.

    Args:
        dirs: List of directories where documents can be accessed
    """
    global ALLOWED_DOCUMENT_DIRS
    ALLOWED_DOCUMENT_DIRS = [os.path.expanduser(d) for d in dirs]
    logger.info("Updated allowed document directories", dirs=ALLOWED_DOCUMENT_DIRS)


def _extract_pdf_text(file_path: Path, page_range: Optional[tuple[int, int]] = None) -> tuple[str, dict]:
    """
    Extract text from a PDF file.

    Args:
        file_path: Path to the PDF file
        page_range: Optional tuple of (start_page, end_page) 1-indexed

    Returns:
        Tuple of (extracted_text, metadata)
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")

    reader = PdfReader(str(file_path))
    num_pages = len(reader.pages)

    # Determine page range
    start_page = 0
    end_page = min(num_pages, MAX_PAGES)

    if page_range:
        start_page = max(0, page_range[0] - 1)  # Convert to 0-indexed
        end_page = min(num_pages, page_range[1])

    # Extract metadata
    metadata = {
        "title": reader.metadata.title if reader.metadata else None,
        "author": reader.metadata.author if reader.metadata else None,
        "subject": reader.metadata.subject if reader.metadata else None,
        "creator": reader.metadata.creator if reader.metadata else None,
        "total_pages": num_pages,
        "pages_extracted": f"{start_page + 1}-{end_page}",
    }

    # Extract text from pages
    text_parts = []
    for i in range(start_page, end_page):
        try:
            page = reader.pages[i]
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
        except Exception as e:
            text_parts.append(f"--- Page {i + 1} ---\n[Error extracting page: {e}]")

    return "\n\n".join(text_parts), metadata


def _extract_docx_text(file_path: Path) -> tuple[str, dict]:
    """
    Extract text from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Tuple of (extracted_text, metadata)
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

    doc = Document(str(file_path))

    # Extract core properties (metadata)
    core_props = doc.core_properties
    metadata = {
        "title": core_props.title,
        "author": core_props.author,
        "subject": core_props.subject,
        "created": str(core_props.created) if core_props.created else None,
        "modified": str(core_props.modified) if core_props.modified else None,
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
    }

    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # Check for heading style
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                text_parts.append(f"\n{'#' * int(level) if level.isdigit() else '##'} {para.text}\n")
            else:
                text_parts.append(para.text)

    # Extract tables
    for i, table in enumerate(doc.tables):
        table_text = [f"\n--- Table {i + 1} ---"]
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            table_text.append(row_text)
        text_parts.append("\n".join(table_text))

    return "\n\n".join(text_parts), metadata


@tool
def document_parser(
    file_path: str,
    page_range: Optional[str] = None,
    include_metadata: bool = True
) -> str:
    """
    Parse and extract text from PDF or DOCX documents.

    Extracts text content while preserving structure where possible.
    For PDFs, can specify a page range. Maximum file size is 50 MB.

    Args:
        file_path: Path to the document file (.pdf or .docx)
        page_range: For PDFs only - page range as "start-end" (e.g., "1-10"). 1-indexed.
        include_metadata: Whether to include document metadata (default: True)

    Returns:
        Extracted text content with optional metadata, or an error message

    Examples:
        - document_parser("/home/user/research_data/report.pdf")
        - document_parser("~/Documents/paper.pdf", page_range="1-5")
        - document_parser("/tmp/research_agent/analysis.docx", include_metadata=False)
    """
    logger.info("Document parse requested", path=file_path)

    # Validate path safety
    is_safe, error = _is_document_path_safe(file_path)
    if not is_safe:
        logger.warning("Document access blocked", path=file_path, reason=error)
        return f"Error: {error}"

    resolved_path = Path(file_path).expanduser().resolve()

    # Check extension
    extension = resolved_path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        return f"Error: Unsupported file type '{extension}'. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"

    # Check if file exists
    if not resolved_path.exists():
        return f"Error: File not found: {file_path}"

    if not resolved_path.is_file():
        return f"Error: Path is not a file: {file_path}"

    # Check file size
    file_size = resolved_path.stat().st_size
    if file_size > MAX_FILE_SIZE:
        return f"Error: File too large ({file_size:,} bytes). Maximum: {MAX_FILE_SIZE:,} bytes"

    # Parse page range for PDFs
    parsed_range = None
    if page_range and extension == ".pdf":
        try:
            parts = page_range.split("-")
            if len(parts) == 2:
                parsed_range = (int(parts[0]), int(parts[1]))
            else:
                return "Error: page_range must be in format 'start-end' (e.g., '1-10')"
        except ValueError:
            return "Error: page_range must contain valid integers"

    # Extract based on file type
    try:
        if extension == ".pdf":
            text, metadata = _extract_pdf_text(resolved_path, parsed_range)
        elif extension in (".docx", ".doc"):
            if extension == ".doc":
                return "Error: .doc format not supported. Please convert to .docx"
            text, metadata = _extract_docx_text(resolved_path)
        else:
            return f"Error: Unsupported file type: {extension}"

        # Build output
        output_parts = []

        if include_metadata and metadata:
            output_parts.append("=== Document Metadata ===")
            for key, value in metadata.items():
                if value:
                    output_parts.append(f"{key}: {value}")
            output_parts.append("\n=== Document Content ===\n")

        output_parts.append(text)

        result = "\n".join(output_parts)

        # Truncate if too long
        if len(result) > MAX_OUTPUT_LENGTH:
            result = result[:MAX_OUTPUT_LENGTH] + f"\n\n... [Content truncated at {MAX_OUTPUT_LENGTH:,} characters]"

        logger.info("Document parsed successfully", path=file_path, length=len(result))
        return result

    except ImportError as e:
        return f"Error: {e}"
    except Exception as e:
        logger.error("Document parsing failed", path=file_path, error=str(e))
        return f"Error parsing document: {str(e)}"


@tool
def document_metadata(file_path: str) -> str:
    """
    Extract only metadata from a PDF or DOCX document.

    Retrieves document properties without extracting full text content.
    Faster than full parsing for large documents.

    Args:
        file_path: Path to the document file (.pdf or .docx)

    Returns:
        Document metadata as formatted text, or an error message

    Examples:
        - document_metadata("/home/user/research_data/report.pdf")
        - document_metadata("~/Documents/paper.docx")
    """
    logger.info("Document metadata requested", path=file_path)

    # Validate path safety
    is_safe, error = _is_document_path_safe(file_path)
    if not is_safe:
        logger.warning("Document access blocked", path=file_path, reason=error)
        return f"Error: {error}"

    resolved_path = Path(file_path).expanduser().resolve()

    # Check extension
    extension = resolved_path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        return f"Error: Unsupported file type '{extension}'. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"

    # Check if file exists
    if not resolved_path.exists():
        return f"Error: File not found: {file_path}"

    # Get file stats
    stat = resolved_path.stat()
    file_size = stat.st_size

    try:
        metadata = {
            "file_name": resolved_path.name,
            "file_size": f"{file_size:,} bytes ({file_size / 1024 / 1024:.2f} MB)",
            "file_type": extension,
        }

        if extension == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(resolved_path))

                metadata["total_pages"] = len(reader.pages)

                if reader.metadata:
                    if reader.metadata.title:
                        metadata["title"] = reader.metadata.title
                    if reader.metadata.author:
                        metadata["author"] = reader.metadata.author
                    if reader.metadata.subject:
                        metadata["subject"] = reader.metadata.subject
                    if reader.metadata.creator:
                        metadata["creator"] = reader.metadata.creator
                    if reader.metadata.producer:
                        metadata["producer"] = reader.metadata.producer
                    if reader.metadata.creation_date:
                        metadata["creation_date"] = str(reader.metadata.creation_date)
                    if reader.metadata.modification_date:
                        metadata["modification_date"] = str(reader.metadata.modification_date)

            except ImportError:
                return "Error: pypdf is required for PDF parsing. Install with: pip install pypdf"

        elif extension == ".docx":
            try:
                from docx import Document
                doc = Document(str(resolved_path))

                props = doc.core_properties
                if props.title:
                    metadata["title"] = props.title
                if props.author:
                    metadata["author"] = props.author
                if props.subject:
                    metadata["subject"] = props.subject
                if props.created:
                    metadata["created"] = str(props.created)
                if props.modified:
                    metadata["modified"] = str(props.modified)
                if props.last_modified_by:
                    metadata["last_modified_by"] = props.last_modified_by

                metadata["total_paragraphs"] = len(doc.paragraphs)
                metadata["total_tables"] = len(doc.tables)

            except ImportError:
                return "Error: python-docx is required for DOCX parsing. Install with: pip install python-docx"

        # Format output
        output_lines = [f"=== Metadata for {resolved_path.name} ===\n"]
        for key, value in metadata.items():
            output_lines.append(f"{key}: {value}")

        logger.info("Metadata extracted successfully", path=file_path)
        return "\n".join(output_lines)

    except Exception as e:
        logger.error("Metadata extraction failed", path=file_path, error=str(e))
        return f"Error extracting metadata: {str(e)}"
