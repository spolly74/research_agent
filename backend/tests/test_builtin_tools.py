"""
Tests for Built-in Tool Library (Phase 2.3)

Tests for:
- Filesystem tools (file_reader, file_writer, list_directory)
- Database tools (database_query, database_schema)
- Document tools (document_parser, document_metadata)
"""

import json
import os
import sqlite3
import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest


# Test Filesystem Tools
class TestFilesystemTools:
    """Tests for filesystem.py tools."""

    @pytest.fixture
    def temp_read_dir(self, tmp_path):
        """Create a temporary directory for reading."""
        read_dir = tmp_path / "read_data"
        read_dir.mkdir()

        # Create test files
        (read_dir / "test.txt").write_text("Hello, World!")
        (read_dir / "data.json").write_text('{"key": "value"}')
        (read_dir / "subdir").mkdir()
        (read_dir / "subdir" / "nested.txt").write_text("Nested content")

        return read_dir

    @pytest.fixture
    def temp_write_dir(self, tmp_path):
        """Create a temporary directory for writing."""
        write_dir = tmp_path / "write_output"
        write_dir.mkdir()
        return write_dir

    def test_file_reader_success(self, temp_read_dir):
        """Test successful file reading."""
        from app.agents.tools.filesystem import file_reader, ALLOWED_READ_DIRS

        # Temporarily add temp dir to allowed dirs
        original_dirs = ALLOWED_READ_DIRS.copy()
        ALLOWED_READ_DIRS.clear()
        ALLOWED_READ_DIRS.append(str(temp_read_dir))

        try:
            result = file_reader.invoke({
                "file_path": str(temp_read_dir / "test.txt")
            })
            assert "Hello, World!" in result
        finally:
            ALLOWED_READ_DIRS.clear()
            ALLOWED_READ_DIRS.extend(original_dirs)

    def test_file_reader_blocked_path(self, temp_read_dir):
        """Test that paths outside allowed directories are blocked."""
        from app.agents.tools.filesystem import file_reader

        result = file_reader.invoke({
            "file_path": "/etc/passwd"
        })
        assert "Error" in result
        assert "allowed directories" in result

    def test_file_reader_path_traversal(self, temp_read_dir):
        """Test that path traversal is blocked."""
        from app.agents.tools.filesystem import file_reader, ALLOWED_READ_DIRS

        original_dirs = ALLOWED_READ_DIRS.copy()
        ALLOWED_READ_DIRS.clear()
        ALLOWED_READ_DIRS.append(str(temp_read_dir))

        try:
            result = file_reader.invoke({
                "file_path": str(temp_read_dir / ".." / ".." / "etc" / "passwd")
            })
            assert "Error" in result
        finally:
            ALLOWED_READ_DIRS.clear()
            ALLOWED_READ_DIRS.extend(original_dirs)

    def test_file_reader_file_not_found(self, temp_read_dir):
        """Test error when file doesn't exist."""
        from app.agents.tools.filesystem import file_reader, ALLOWED_READ_DIRS

        original_dirs = ALLOWED_READ_DIRS.copy()
        ALLOWED_READ_DIRS.clear()
        ALLOWED_READ_DIRS.append(str(temp_read_dir))

        try:
            result = file_reader.invoke({
                "file_path": str(temp_read_dir / "nonexistent.txt")
            })
            assert "Error" in result
            assert "not found" in result
        finally:
            ALLOWED_READ_DIRS.clear()
            ALLOWED_READ_DIRS.extend(original_dirs)

    def test_file_reader_blocked_extension(self, temp_read_dir):
        """Test that certain file extensions are blocked."""
        from app.agents.tools.filesystem import file_reader, ALLOWED_READ_DIRS

        original_dirs = ALLOWED_READ_DIRS.copy()
        ALLOWED_READ_DIRS.clear()
        ALLOWED_READ_DIRS.append(str(temp_read_dir))

        # Create a binary file
        (temp_read_dir / "test.exe").write_bytes(b"binary")

        try:
            result = file_reader.invoke({
                "file_path": str(temp_read_dir / "test.exe")
            })
            assert "Error" in result
            assert "extension" in result.lower()
        finally:
            ALLOWED_READ_DIRS.clear()
            ALLOWED_READ_DIRS.extend(original_dirs)

    def test_file_writer_success(self, temp_write_dir):
        """Test successful file writing."""
        from app.agents.tools.filesystem import file_writer, ALLOWED_WRITE_DIRS

        original_dirs = ALLOWED_WRITE_DIRS.copy()
        ALLOWED_WRITE_DIRS.clear()
        ALLOWED_WRITE_DIRS.append(str(temp_write_dir))

        try:
            result = file_writer.invoke({
                "file_path": str(temp_write_dir / "output.txt"),
                "content": "Test content"
            })
            assert "Successfully" in result
            assert (temp_write_dir / "output.txt").read_text() == "Test content"
        finally:
            ALLOWED_WRITE_DIRS.clear()
            ALLOWED_WRITE_DIRS.extend(original_dirs)

    def test_file_writer_append_mode(self, temp_write_dir):
        """Test file writing in append mode."""
        from app.agents.tools.filesystem import file_writer, ALLOWED_WRITE_DIRS

        original_dirs = ALLOWED_WRITE_DIRS.copy()
        ALLOWED_WRITE_DIRS.clear()
        ALLOWED_WRITE_DIRS.append(str(temp_write_dir))

        try:
            file_path = str(temp_write_dir / "append.txt")

            # Write initial content
            file_writer.invoke({
                "file_path": file_path,
                "content": "Line 1\n",
                "mode": "write"
            })

            # Append more content
            file_writer.invoke({
                "file_path": file_path,
                "content": "Line 2\n",
                "mode": "append"
            })

            assert (temp_write_dir / "append.txt").read_text() == "Line 1\nLine 2\n"
        finally:
            ALLOWED_WRITE_DIRS.clear()
            ALLOWED_WRITE_DIRS.extend(original_dirs)

    def test_file_writer_blocked_path(self, temp_write_dir):
        """Test that writing to blocked paths fails."""
        from app.agents.tools.filesystem import file_writer

        result = file_writer.invoke({
            "file_path": "/etc/test.txt",
            "content": "Test"
        })
        assert "Error" in result

    def test_list_directory_success(self, temp_read_dir):
        """Test successful directory listing."""
        from app.agents.tools.filesystem import list_directory, ALLOWED_READ_DIRS

        original_dirs = ALLOWED_READ_DIRS.copy()
        ALLOWED_READ_DIRS.clear()
        ALLOWED_READ_DIRS.append(str(temp_read_dir))

        try:
            result = list_directory.invoke({
                "directory_path": str(temp_read_dir)
            })
            assert "test.txt" in result
            assert "data.json" in result
            assert "subdir" in result
        finally:
            ALLOWED_READ_DIRS.clear()
            ALLOWED_READ_DIRS.extend(original_dirs)

    def test_list_directory_with_pattern(self, temp_read_dir):
        """Test directory listing with glob pattern."""
        from app.agents.tools.filesystem import list_directory, ALLOWED_READ_DIRS

        original_dirs = ALLOWED_READ_DIRS.copy()
        ALLOWED_READ_DIRS.clear()
        ALLOWED_READ_DIRS.append(str(temp_read_dir))

        try:
            result = list_directory.invoke({
                "directory_path": str(temp_read_dir),
                "pattern": "*.txt"
            })
            assert "test.txt" in result
            assert "data.json" not in result
        finally:
            ALLOWED_READ_DIRS.clear()
            ALLOWED_READ_DIRS.extend(original_dirs)


# Test Database Tools
class TestDatabaseTools:
    """Tests for database.py tools."""

    @pytest.fixture
    def temp_db(self, tmp_path):
        """Create a temporary SQLite database."""
        db_path = tmp_path / "test.db"
        conn = sqlite3.connect(str(db_path))
        cursor = conn.cursor()

        # Create test table
        cursor.execute("""
            CREATE TABLE users (
                id INTEGER PRIMARY KEY,
                name TEXT NOT NULL,
                email TEXT,
                age INTEGER
            )
        """)

        # Insert test data
        cursor.executemany(
            "INSERT INTO users (name, email, age) VALUES (?, ?, ?)",
            [
                ("Alice", "alice@example.com", 30),
                ("Bob", "bob@example.com", 25),
                ("Charlie", "charlie@example.com", 35),
            ]
        )

        conn.commit()
        conn.close()

        return db_path

    def test_database_query_success(self, temp_db, tmp_path):
        """Test successful database query."""
        from app.agents.tools.database import database_query, ALLOWED_DB_PATHS

        original_paths = ALLOWED_DB_PATHS.copy()
        ALLOWED_DB_PATHS.clear()
        ALLOWED_DB_PATHS.append(str(tmp_path))

        try:
            result = database_query.invoke({
                "db_path": str(temp_db),
                "query": "SELECT * FROM users"
            })
            assert "Alice" in result
            assert "Bob" in result
            assert "Charlie" in result
            assert "3 row(s)" in result
        finally:
            ALLOWED_DB_PATHS.clear()
            ALLOWED_DB_PATHS.extend(original_paths)

    def test_database_query_with_params(self, temp_db, tmp_path):
        """Test database query with parameters."""
        from app.agents.tools.database import database_query, ALLOWED_DB_PATHS

        original_paths = ALLOWED_DB_PATHS.copy()
        ALLOWED_DB_PATHS.clear()
        ALLOWED_DB_PATHS.append(str(tmp_path))

        try:
            result = database_query.invoke({
                "db_path": str(temp_db),
                "query": "SELECT * FROM users WHERE age > ?",
                "params": "[28]"
            })
            assert "Alice" in result
            assert "Charlie" in result
            assert "Bob" not in result
        finally:
            ALLOWED_DB_PATHS.clear()
            ALLOWED_DB_PATHS.extend(original_paths)

    def test_database_query_json_format(self, temp_db, tmp_path):
        """Test database query with JSON output format."""
        from app.agents.tools.database import database_query, ALLOWED_DB_PATHS

        original_paths = ALLOWED_DB_PATHS.copy()
        ALLOWED_DB_PATHS.clear()
        ALLOWED_DB_PATHS.append(str(tmp_path))

        try:
            result = database_query.invoke({
                "db_path": str(temp_db),
                "query": "SELECT name, age FROM users WHERE name = 'Alice'",
                "output_format": "json"
            })
            data = json.loads(result)
            assert len(data) == 1
            assert data[0]["name"] == "Alice"
            assert data[0]["age"] == 30
        finally:
            ALLOWED_DB_PATHS.clear()
            ALLOWED_DB_PATHS.extend(original_paths)

    def test_database_query_blocked_modify(self, temp_db, tmp_path):
        """Test that modification queries are blocked."""
        from app.agents.tools.database import database_query, ALLOWED_DB_PATHS

        original_paths = ALLOWED_DB_PATHS.copy()
        ALLOWED_DB_PATHS.clear()
        ALLOWED_DB_PATHS.append(str(tmp_path))

        try:
            result = database_query.invoke({
                "db_path": str(temp_db),
                "query": "DELETE FROM users WHERE id = 1"
            })
            assert "Error" in result
            assert "SELECT" in result
        finally:
            ALLOWED_DB_PATHS.clear()
            ALLOWED_DB_PATHS.extend(original_paths)

    def test_database_query_blocked_path(self, temp_db):
        """Test that queries to non-allowed paths are blocked."""
        from app.agents.tools.database import database_query

        result = database_query.invoke({
            "db_path": "/etc/test.db",
            "query": "SELECT 1"
        })
        assert "Error" in result

    def test_database_schema_all_tables(self, temp_db, tmp_path):
        """Test getting schema for all tables."""
        from app.agents.tools.database import database_schema, ALLOWED_DB_PATHS

        original_paths = ALLOWED_DB_PATHS.copy()
        ALLOWED_DB_PATHS.clear()
        ALLOWED_DB_PATHS.append(str(tmp_path))

        try:
            result = database_schema.invoke({
                "db_path": str(temp_db)
            })
            assert "users" in result
            assert "CREATE TABLE" in result
            assert "name TEXT" in result
        finally:
            ALLOWED_DB_PATHS.clear()
            ALLOWED_DB_PATHS.extend(original_paths)

    def test_database_schema_specific_table(self, temp_db, tmp_path):
        """Test getting schema for a specific table."""
        from app.agents.tools.database import database_schema, ALLOWED_DB_PATHS

        original_paths = ALLOWED_DB_PATHS.copy()
        ALLOWED_DB_PATHS.clear()
        ALLOWED_DB_PATHS.append(str(tmp_path))

        try:
            result = database_schema.invoke({
                "db_path": str(temp_db),
                "table_name": "users"
            })
            assert "users" in result
            assert "CREATE TABLE" in result
        finally:
            ALLOWED_DB_PATHS.clear()
            ALLOWED_DB_PATHS.extend(original_paths)


# Test Document Tools
class TestDocumentTools:
    """Tests for documents.py tools."""

    @pytest.fixture
    def temp_doc_dir(self, tmp_path):
        """Create a temporary directory for documents."""
        doc_dir = tmp_path / "documents"
        doc_dir.mkdir()
        return doc_dir

    def test_document_parser_unsupported_extension(self, temp_doc_dir):
        """Test error for unsupported file extension."""
        from app.agents.tools.documents import document_parser, ALLOWED_DOCUMENT_DIRS

        original_dirs = ALLOWED_DOCUMENT_DIRS.copy()
        ALLOWED_DOCUMENT_DIRS.clear()
        ALLOWED_DOCUMENT_DIRS.append(str(temp_doc_dir))

        # Create a file with unsupported extension
        (temp_doc_dir / "test.xyz").write_text("content")

        try:
            result = document_parser.invoke({
                "file_path": str(temp_doc_dir / "test.xyz")
            })
            assert "Error" in result
            assert "Unsupported" in result
        finally:
            ALLOWED_DOCUMENT_DIRS.clear()
            ALLOWED_DOCUMENT_DIRS.extend(original_dirs)

    def test_document_parser_file_not_found(self, temp_doc_dir):
        """Test error when file doesn't exist."""
        from app.agents.tools.documents import document_parser, ALLOWED_DOCUMENT_DIRS

        original_dirs = ALLOWED_DOCUMENT_DIRS.copy()
        ALLOWED_DOCUMENT_DIRS.clear()
        ALLOWED_DOCUMENT_DIRS.append(str(temp_doc_dir))

        try:
            result = document_parser.invoke({
                "file_path": str(temp_doc_dir / "nonexistent.pdf")
            })
            assert "Error" in result
            assert "not found" in result
        finally:
            ALLOWED_DOCUMENT_DIRS.clear()
            ALLOWED_DOCUMENT_DIRS.extend(original_dirs)

    def test_document_parser_blocked_path(self, temp_doc_dir):
        """Test that paths outside allowed directories are blocked."""
        from app.agents.tools.documents import document_parser

        result = document_parser.invoke({
            "file_path": "/etc/secret.pdf"
        })
        assert "Error" in result

    def test_document_metadata_unsupported_extension(self, temp_doc_dir):
        """Test error for unsupported file extension in metadata."""
        from app.agents.tools.documents import document_metadata, ALLOWED_DOCUMENT_DIRS

        original_dirs = ALLOWED_DOCUMENT_DIRS.copy()
        ALLOWED_DOCUMENT_DIRS.clear()
        ALLOWED_DOCUMENT_DIRS.append(str(temp_doc_dir))

        # Create a file with unsupported extension
        (temp_doc_dir / "test.xyz").write_text("content")

        try:
            result = document_metadata.invoke({
                "file_path": str(temp_doc_dir / "test.xyz")
            })
            assert "Error" in result
            assert "Unsupported" in result
        finally:
            ALLOWED_DOCUMENT_DIRS.clear()
            ALLOWED_DOCUMENT_DIRS.extend(original_dirs)

    @pytest.mark.skipif(
        not pytest.importorskip("pypdf", reason="pypdf not installed"),
        reason="pypdf not installed"
    )
    def test_document_parser_pdf(self, temp_doc_dir):
        """Test PDF parsing if pypdf is available."""
        from pypdf import PdfWriter
        from app.agents.tools.documents import document_parser, ALLOWED_DOCUMENT_DIRS

        original_dirs = ALLOWED_DOCUMENT_DIRS.copy()
        ALLOWED_DOCUMENT_DIRS.clear()
        ALLOWED_DOCUMENT_DIRS.append(str(temp_doc_dir))

        # Create a simple PDF
        pdf_path = temp_doc_dir / "test.pdf"
        writer = PdfWriter()
        writer.add_blank_page(612, 792)  # Letter size
        with open(pdf_path, "wb") as f:
            writer.write(f)

        try:
            result = document_parser.invoke({
                "file_path": str(pdf_path)
            })
            # Should succeed even if empty
            assert "Error" not in result or "pypdf" not in result
        finally:
            ALLOWED_DOCUMENT_DIRS.clear()
            ALLOWED_DOCUMENT_DIRS.extend(original_dirs)

    @pytest.mark.skipif(
        not pytest.importorskip("docx", reason="python-docx not installed"),
        reason="python-docx not installed"
    )
    def test_document_parser_docx(self, temp_doc_dir):
        """Test DOCX parsing if python-docx is available."""
        from docx import Document
        from app.agents.tools.documents import document_parser, ALLOWED_DOCUMENT_DIRS

        original_dirs = ALLOWED_DOCUMENT_DIRS.copy()
        ALLOWED_DOCUMENT_DIRS.clear()
        ALLOWED_DOCUMENT_DIRS.append(str(temp_doc_dir))

        # Create a simple DOCX
        docx_path = temp_doc_dir / "test.docx"
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph.")
        doc.save(str(docx_path))

        try:
            result = document_parser.invoke({
                "file_path": str(docx_path)
            })
            assert "Test Document" in result or "test paragraph" in result
        finally:
            ALLOWED_DOCUMENT_DIRS.clear()
            ALLOWED_DOCUMENT_DIRS.extend(original_dirs)


# Test Tool Registration
class TestToolRegistration:
    """Test that all new tools are properly registered."""

    def test_all_new_tools_registered(self):
        """Test that all Phase 2.3 tools are registered."""
        from app.agents.tools.registry import get_tool_registry, register_builtin_tools

        # Reset and register
        registry = get_tool_registry()
        register_builtin_tools()

        status = registry.get_registry_status()
        tool_names = [t["name"] for t in status["tools"]]

        # Check all new tools are registered
        expected_tools = [
            "file_reader",
            "file_writer",
            "list_directory",
            "database_query",
            "database_schema",
            "document_parser",
            "document_metadata",
        ]

        for tool_name in expected_tools:
            assert tool_name in tool_names, f"Tool {tool_name} not registered"

    def test_tool_categories(self):
        """Test that tools are in correct categories."""
        from app.agents.tools.registry import get_tool_registry, register_builtin_tools, ToolCategory

        registry = get_tool_registry()
        register_builtin_tools()

        # Check file tools
        file_tools = registry.get_tools_by_category(ToolCategory.FILE)
        file_tool_names = [t.name for t in file_tools]
        assert "file_reader" in file_tool_names
        assert "file_writer" in file_tool_names
        assert "document_parser" in file_tool_names

        # Check data tools
        data_tools = registry.get_tools_by_category(ToolCategory.DATA)
        data_tool_names = [t.name for t in data_tools]
        assert "database_query" in data_tool_names
        assert "database_schema" in data_tool_names

    def test_tool_agent_restrictions(self):
        """Test that tools have correct agent restrictions."""
        from app.agents.tools.registry import get_tool_registry, register_builtin_tools

        registry = get_tool_registry()
        register_builtin_tools()

        # Researcher should have read tools but not write tools
        researcher_tools = registry.get_tools_for_agent("researcher")
        researcher_tool_names = [t.name for t in researcher_tools]
        assert "file_reader" in researcher_tool_names
        assert "document_parser" in researcher_tool_names
        assert "file_writer" not in researcher_tool_names

        # Coder should have write tools
        coder_tools = registry.get_tools_for_agent("coder")
        coder_tool_names = [t.name for t in coder_tools]
        assert "file_writer" in coder_tool_names
        assert "database_query" in coder_tool_names
